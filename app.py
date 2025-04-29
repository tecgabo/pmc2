# Project Model Canvas - PRO Edition (Exportação Completa)
import streamlit as st
from PIL import Image, ImageDraw, ImageFont
from fpdf import FPDF
from docx import Document
import datetime
import io

# Configurações
st.set_page_config(
    page_title="Project Model Canvas - PRO Edition",
    page_icon="📓",
    layout="wide"
)

# Inicializar dados
if 'dados' not in st.session_state:
    st.session_state.dados = {campo: '' for campo in [
        'nome_projeto', 'responsavel', 'data',
        'justificativas', 'pitch', 'produto', 'stakeholders', 'premissas', 'riscos',
        'objetivos', 'requisitos', 'equipe', 'entregas', 'cronograma', 'beneficios',
        'restricoes', 'custos', 'observacoes']}
    st.session_state.dados['data'] = datetime.date.today().strftime("%d/%m/%Y")

# Funções auxiliares
def gerar_imagem_canvas(data):
    img = Image.new('RGB', (1800, 1200), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 24)
    except:
        font = ImageFont.load_default()
    y = 50
    for campo, texto in data.items():
        draw.text((50, y), f"{campo.upper()}: {texto}", fill=(0, 0, 0), font=font)
        y += 70
    return img

def gerar_pdf_resumo(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Resumo do Projeto - Project Model Canvas", ln=True, align="C")
    pdf.ln(10)
    for campo, texto in data.items():
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, campo.capitalize(), ln=True)
        pdf.set_font("Arial", '', 12)
        pdf.multi_cell(0, 10, texto)
        pdf.ln(2)
    buffer = io.BytesIO()
    pdf.output(buffer)
    return buffer.getvalue()

def gerar_docx_termo_abertura(data):
    doc = Document()
    doc.add_heading('Termo de Abertura do Projeto', 0)
    doc.add_paragraph(f"Nome do Projeto: {data.get('nome_projeto', '')}")
    doc.add_paragraph(f"Responsável: {data.get('responsavel', '')}")
    doc.add_paragraph(f"Data: {data.get('data', '')}")
    doc.add_heading('Objetivos', level=1)
    doc.add_paragraph(data.get('objetivos', ''))
    doc.add_heading('Justificativas', level=1)
    doc.add_paragraph(data.get('justificativas', ''))
    doc.add_heading('Principais Entregas', level=1)
    doc.add_paragraph(data.get('entregas', ''))
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Sidebar
st.sidebar.title("Menu")
menu = st.sidebar.radio(
    "Navegar:",
    ("Início", "Preencher Canvas", "Exportar Projeto")
)

# Conteúdo
if menu == "Início":
    st.title("📓 Project Model Canvas - PRO Edition")
    st.markdown("Ferramenta completa para estruturar, planejar e exportar projetos.")

elif menu == "Preencher Canvas":
    st.title("📝 Preenchimento do Project Model Canvas")
    st.session_state.dados['nome_projeto'] = st.text_input("Nome do Projeto:", value=st.session_state.dados['nome_projeto'])
    st.session_state.dados['responsavel'] = st.text_input("Responsável:", value=st.session_state.dados['responsavel'])
    st.session_state.dados['data'] = st.date_input("Data:", datetime.datetime.strptime(st.session_state.dados['data'], "%d/%m/%Y")).strftime("%d/%m/%Y")
    
    st.subheader("Áreas do Canvas")
    for campo in ['justificativas', 'pitch', 'produto', 'stakeholders', 'premissas', 'riscos',
                  'objetivos', 'requisitos', 'equipe', 'entregas', 'cronograma', 'beneficios',
                  'restricoes', 'custos', 'observacoes']:
        st.session_state.dados[campo] = st.text_area(campo.capitalize(), value=st.session_state.dados[campo])

elif menu == "Exportar Projeto":
    st.title("📥 Exportar Projeto Completo")

    if st.button("Gerar Imagem do Canvas (.PNG)"):
        img = gerar_imagem_canvas(st.session_state.dados)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        st.download_button(
            label="Baixar Canvas PNG",
            data=buf.getvalue(),
            file_name="canvas_projeto.png",
            mime="image/png"
        )

    if st.button("Gerar Resumo Executivo (.PDF)"):
        pdf = gerar_pdf_resumo(st.session_state.dados)
        st.download_button(
            label="Baixar Resumo PDF",
            data=pdf,
            file_name="resumo_projeto.pdf",
            mime="application/pdf"
        )

    if st.button("Gerar Termo de Abertura (.DOCX)"):
        docx = gerar_docx_termo_abertura(st.session_state.dados)
        st.download_button(
            label="Baixar Termo de Abertura DOCX",
            data=docx,
            file_name="termo_abertura.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
